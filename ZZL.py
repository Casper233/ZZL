import argparse
import os
import shutil
import time
from datetime import datetime

import docx

timestamp = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())

def close_window():
    if os.name == "nt":
        # Windows系统
        os.system("taskkill /f /im python.exe")
    else:
        # 其他系统
        os.system("kill -9 {}".format(os.getpid()))


# LogMaker
def create_log_file():
    folder = "logs"
    #timestamp = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
    if not os.path.exists(folder):
        os.makedirs(folder)
    filename = f"{folder}/{timestamp}.txt"
    with open(filename, "w") as file:
        pass


def logmake(stra, message, level):
    localtimestamp = time.strftime("%Y-%m-%d %H-%M-%S", time.localtime())
    folder = "logs"
    filename = f"{folder}/{timestamp}.txt"
    with open(filename, "a", encoding="utf-8") as file:
        if level == 0:
            file.write("[" + localtimestamp + "] [" + stra + "/INFO] " + message + "\n")
        elif level == 1:
            file.write("[" + localtimestamp + "] [" + stra + "/WARNING] " + message + "\n")
        elif level == 2:
            file.write("[" + localtimestamp + "] [" + stra + "/ERROR] " + message + "\n")
        elif level == 3:
            file.write("[" + localtimestamp + "] [" + stra + "/CRITICAL] " + message + "\n")
        else:
            file.write("[" + localtimestamp + "] [" + stra + "/DEBUG] " + message + "\n")


# Copy File Area
def create_folder(week_num: str, grade_name: str, class_name: str) -> str:
    """创建文件夹，并返回文件夹路径"""
    folder_name = f"第{week_num}周 {grade_name}{class_name}"
    os.mkdir(folder_name)
    logmake("CopyFile_MainThread", f"{datetime.now()} 周数为 {week_num}\n", 0)
    logmake("CopyFile_MainThread", f"{datetime.now()} 班级为 {grade_name}{class_name}\n", 0)
    logmake("CopyFile_MainThread", f"{datetime.now()} 创建文件夹成功\n", 0)
    return folder_name


def copy_template_files(template_dir: str, dest_dir: str) -> None:
    try:
        for file_name in os.listdir(template_dir):
            src_file = os.path.join(template_dir, file_name)
            dest_file = os.path.join(dest_dir, file_name)
            shutil.copy(src_file, dest_file)
        logmake("CopyFile_MainThread", f"{datetime.now()} 复制文件成功\n", 0)
    except Exception as e:
        logmake("CopyFile_MainThread", f"{datetime.now()} 复制文件失败：{e}\n", 1)


def FileCopy():
    week_num = input("请输入周数：")
    grade_list = ["初一", "初二", "初三", "高一", "高二", "高三"]
    grade_num = input(f"请选择年级（输入数字 1-{len(grade_list)}）：\n{grade_list}\n")

    try:
        grade_num = int(grade_num)
        if grade_num < 1 or grade_num > len(grade_list):
            raise ValueError
    except ValueError:
        print("输入不合法，请重新输入")
        logmake("CopyFile_MainThread", f"{datetime.now()} 输入不合法,输入为 {grade_num}", 2)
        return

    grade_name = grade_list[int(grade_num) - 1]
    class_name = input("请输入班级：")
    class_name = f"({class_name})班"
    folder_path = create_folder(week_num, grade_name, class_name)
    template_dir = "模板"

    if not os.path.exists(template_dir):
        os.mkdir(template_dir)
    copy_template_files(template_dir, folder_path)

    return 1


def FileCopy_Start():
    FileCopy()
    if FileCopy() == "None":
        FileCopy()


# Distribute Area


def Input_Data():
    parser = argparse.ArgumentParser(description='处理Word文件中的文本信息')
    parser.add_argument('path', metavar='path', type=str, nargs='?', default=os.getcwd(),
                        help='指定要处理的文件或文件夹路径，默认为当前工作目录')
    args = parser.parse_args()
    # 判断要处理的路径是否存在
    if not os.path.exists(args.path):
        logmake("FileInput_MainThread",f"路径 {args.path} 不存在",1)
        print(f"路径 {args.path} 不存在")
        return
    # 判断要处理的路径是文件还是文件夹
    if os.path.isfile(args.path) and args.path.endswith('.docx'):
        # 处理单个Word文件
        logmake("FileInput_MainThread",f"处理文件 {args.path} 中的文本信息...", 0)
        print(f"处理文件 {args.path} 中的文本信息...")
        deal_with_docx(args.path)
    elif os.path.isdir(args.path):
        # 处理一个文件夹中的所有Word文件
        os.chdir(args.path)  # 将当前工作目录更改为指定路径
        files = list_files(args.path)
        print("目录列表：")
        for i, f in enumerate(files):
            print(f" {i + 1}. {f}")
        deal_with_folder(args.path)
    else:
        logmake("FileInput_MainThread", f"路径 {args.path} 不是一个有效的Word文件或文件夹！", 2)
        print(f"路径 {args.path} 不是一个有效的Word文件或文件夹！")

    '''
    
    folder_path = choose_file_or_folder()
    logmake("FileInput_MainThread", f"{datetime.now()} 输入的文件夹路径为 {folder_path}", 0)
    # 进入文件夹，复制其中的”模板.docx“并重命名
    template_path = os.path.join(folder_path, "x.x 星期x（高中模板）.docx")
    if os.path.exists(template_path):
        now = datetime.now()
        date = now.strftime("%m.%d")
        weekday = now.strftime("%A")
        new_file_name = date + weekday + ".docx"
        new_file_path = os.path.join(folder_path, new_file_name)
        shutil.copy(template_path, new_file_path)
        logmake("FileInput_MainThread", f"{datetime.now()} 复制模板文件成功", 0)
    else:
        print("模板文件不存在，请确保模板文件位于指定的文件夹中")
        logmake("FileInput_MainThread", f"{datetime.now()} 模板文件不存在，请确保模板文件位于指定的文件夹中", 1)
        return
    
    # 读取word文件并进行拆分
    document = Document(new_file_path)
    sections = []
    List_EntT = []
    with open(template_path, 'r', encoding='utf-8') as file:
        for line in file:
            if line.startswith('大标题'):
                continue
            List_EntT.append(line.strip())
    
    for section in document.sections:
        sections.append(section)
    sections.append(None)
    headings = []
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            headings.append(paragraph.text)
    headings.append(None)
    sections_headings = []
    for i in range(len(sections) - 1):
        section_headings = []
        for heading in headings:
            if heading == None or sections[i].start <= heading._element.xpath('ancestor::w:p')[0].getparent().index(
                    sections[i + 1].start):
                break
            section_headings.append(heading)
        sections_headings.append(section_headings)
    
    # 在每个大标题下让用户输入信息
    for i in range(len(sections) - 1):
        print("以下是第", i + 1, "个大标题下的信息：")
        for heading in sections_headings[i]:
            print(heading)
            text = input("请输入该部分的信息：")
            if text == "":
                text = "无情况"
            for paragraph in document.paragraphs:
                if paragraph.style.name.startswith("Heading") and paragraph.text == heading:
                    next_paragraph = paragraph._element.getnext()
                    while next_paragraph is not None and not next_paragraph.xpath('name() = "w:p"'):
                        next_paragraph = next_paragraph.getnext()
                    if next_paragraph is None:
                        document.add_paragraph(text)
                    else:
                        next_paragraph.insert_before(text)
        print()

    # 保存word文件
    document.save(new_file_path)
    '''

def deal_with_folder(foldername):
    # 处理一个文件夹中的所有Word文件
    while True:
        files = os.listdir(foldername)
        docx_files = [f for f in files if os.path.isfile(os.path.join(foldername, f)) and f.endswith('.docx')]
        if len(docx_files) == 0:
            logmake("FileInput_Dealfolder", "文件夹中没有Word文件", 2)
            print("文件夹中没有Word文件")
            foldername, files = choose_file_or_folder(files, foldername)
        else:
            break

    # 让用户选择要处理的 Word 文件
    if len(docx_files) > 1:
        logmake("FileInput_Dealfolder", f"文件夹中有以下 {len(docx_files)} 个Word文件：", 0)
        print(f"文件夹中有以下 {len(docx_files)} 个Word文件：")
        for i, f in enumerate(docx_files):
            logmake("FileInput_Dealfolder", f"  {i+1}. {f}", 0)
        while True:
            choice = input("请选择一个文件（输入数字）：")
            if choice.isdigit() and int(choice) >= 1 and int(choice) <= len(docx_files):
                filename = docx_files[int(choice)-1]
                filepath = os.path.join(foldername, filename)
                logmake("FileInput_Dealfolder", f"处理文件 {filepath} 中的文本信息...", 0)
                print(f"处理文件 {filepath} 中的文本信息...")
                deal_with_docx(filepath)
                break
    else:
        filename = docx_files[0]
        filepath = os.path.join(foldername, filename)
        logmake("FileInput_Dealfolder", f"处理文件 {filepath} 中的文本信息...", 0)
        print(f"处理文件 {filepath} 中的文本信息...")
        deal_with_docx(filepath)

def deal_with_docx(filename):
    try:
        doc = docx.Document(filename)
    except Exception as e:
        logmake("FileInput_Dealdoc", f"文件 {filename} 不是有效的 Word 文件或者无法打开", 2)
        print(f"文件 {filename} 不是有效的 Word 文件或者无法打开")
        choice = input("是否重新选择文件？（y/n）")
        if choice.lower() == 'y':
            return
        else:
            raise e

    # 将文本信息存储到文件中
    with open('dealer_temp/deal.txt', 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            if para.text.strip():  # 只处理非空段落
                f.write(para.text + '\n')
    logmake("FileInput_Dealer", f"文件 {filename} 中的文本信息已经存储到 dealer_temp/deal.txt 文件中", 0)
    print(f"文件 {filename} 中的文本信息已经存储到 dealer_temp/deal.txt 文件中")

def list_files(path, show_files=True):
    logmake("FileInput_Filelist", f"列出目录 {os.path.abspath(path)} 中的文件和文件夹", 0)
    try:
        files = os.listdir(path)
    except Exception as e:
        print(f"{e}")
        return []
    if show_files:
        for i, f in enumerate(files):
            if os.path.isdir(os.path.join(path, f)):
                files[i] = f + '/'
            logmake("FileInput_Filelist", f"{i+1}. {files[i]}", 0)
            print(f" {i+1}. {files[i]}")
    else:
        files = []
    return files


# 选择一个文件或文件夹
def choose_file_or_folder(files, path, show_files=True):
    # 选择一个文件或文件夹
    while True:
        choice = input("请选择一个文件或文件夹（输入数字）或输入'..'(返回上一级目录)：")
        if choice == '..':
            path = os.path.dirname(os.path.abspath(path))  # 返回上一级目录
            files = list_files(path, show_files=True)  # 列出上一级目录的文件和文件夹
        elif choice.isdigit() and int(choice) >= 1 and int(choice) <= len(files):
            filename = files[int(choice)-1]
            if filename.endswith('/'):
                filename = filename[:-1]  # 去掉文件夹名称末尾的 '/'
                path = os.path.join(path, filename)  # 进入选择的文件夹
                files = list_files(path, show_files=True)  # 列出该文件夹中的文件和文件夹
            else:
                path = os.path.join(path, filename)  # 选择的是文件，更新路径
            return path, files


def Main():

    create_log_file()

    logmake("System_Event_Handler", f" 主程序已启动", 0)

    program_list = ["新周复制", "记录输入数据", "统计某日数据整合"]

    while True:
        print("请选择要执行的程序：")
        for i, program in enumerate(program_list):
            print(f"{i + 1}. {program}")
        print("0. 退出程序")

        choice = input("请输入程序编号：")
        if choice == "0":
            logmake("System_Event_Handler", f" 选择的程序编号为 {choice}", 0)
            logmake("System_Event_Handler", f" 正在关闭程序", 0)
            logmake("System_Event_Handler", f" Bye~", 0)
            break

        try:
            choice = int(choice)
            if choice < 1 or choice > len(program_list):
                raise ValueError
        except ValueError:
            print("输入不合法，请重新输入")
            continue
        logmake("System_Event_Handler", f"{datetime.now()} 选择的程序编号为 {choice}", 0)

        program_name = program_list[choice - 1]
        logmake("System_Event_Handler", f"{datetime.now()} 选择的程序为 {program_name} ", 0)
        print(f"正在打开 {program_name} ...")
        # 在这里执行对应的程序
        if program_name == "新周复制":
            FileCopy_Start()
            print("程序已完成")
        elif program_name == "记录输入数据":
            Input_Data()
            print("程序已完成")
        elif program_name == "统计某日数据整合":
            print("程序已完成")

    print("程序已退出")


if __name__ == "__main__":
    Main()
